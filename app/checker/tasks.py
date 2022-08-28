import docx2txt
import requests
from celery import shared_task
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from requests.exceptions import InvalidJSONError

from checker.models import Paragraph, Docx, WordDocx, WordParagraph
from checker.services.file import process_paragraphs, process_word_paragraphs, split_text


@shared_task()
def process_file(pk: int, *args, **kwargs):
    file = Docx.objects.get(pk=pk)
    uuid = file.uuid
    document = docx2txt.process(file.file.path)
    # paragraphs = process_paragraphs(document)
    paragraphs, groups = split_text(document)

    file.paragraphs_loaded = len(paragraphs)
    file.save(update_fields=["paragraphs_loaded"])

    cut = 10
    for i in range(len(paragraphs) // cut):
        vals = [x for x in range(i * cut, (i+ 1) * cut)]
        dct = {x: paragraphs[x] for x in vals}
        x = requests.post("http://109.248.175.223:5000/api", json=dct, timeout=1)
        for el_id, dat in x.json().items():
            type_id, score = dat
            Paragraph.objects.create(
                type_id=type_id, docx=file, text=str(groups[int(el_id)]) + dct[int(el_id)] + str(groups[int(el_id)]), score=score
            )

    #for i in range(0, len(paragraphs) // cut + 1):
    #    vals = paragraphs[i * cut : (i + 1) * cut + 1]
    #    dct = {x: vals[x] for x in range(len(vals))}
    #
    #    x = requests.post("http://109.248.175.223:5000/api", json=dct)
    #    if x.status_code == 200:
    #        try:
    #            for el_id, dat in x.json().items():
    #                type_id, score = dat
    #                Paragraph.objects.create(
    #                    type_id=type_id, docx=file, text=str(groups[g_c])+dct[int(el_id)]+str(groups[g_c]), score=score
    #                )
    #                g_c += 1
    #
    #            counter += len(vals)
    #            print(f"processing {uuid}, {counter}/{len_c}")
    #            file.paragraphs_processed = counter
    #            file.save(update_fields=["paragraphs_processed"])
    #        except InvalidJSONError:
    #            print("json pars error")
    #    else:
    #        print(f"AI server error, {x.status_code}")


    return pk


@shared_task()
def process_word(pk: int, *args, **kwargs):
    file = WordDocx.objects.get(pk=pk)
    uuid = file.uuid
    paragraphs = process_word_paragraphs(file.text.tobytes().decode())

    file.paragraphs_loaded = len(paragraphs)
    file.save(update_fields=["paragraphs_loaded"])

    cut = 10
    len_c = len(paragraphs) + 1
    paragraphs = list(paragraphs.values())
    counter = 0
    for i in range(0, len(paragraphs) // cut + 1):
        vals = paragraphs[i * cut : (i + 1) * cut + 1]
        dct = {x: vals[x] for x in range(len(vals))}

        x = requests.post("http://109.248.175.223:5000/api", json=dct)
        if x.status_code == 200:
            try:
                for el_id, dat in x.json().items():
                    type_id, score = dat
                    WordParagraph.objects.create(
                        type_id=type_id, docx=file, text=dct[int(el_id)], score=score
                    )

                counter += len(vals)
                print(f"processing {uuid}, {counter}/{len_c}")
                file.paragraphs_processed = counter
                file.save(update_fields=["paragraphs_processed"])
            except InvalidJSONError:
                print("json pars error")
        else:
            print(f"AI server error, {x.status_code}")

    return pk


@shared_task
def highlight_file(pk: int, *args, **kwargs):
    c = 0
    lim = 0
    file = Docx.objects.get(pk=pk)
    document = Document(file.file.path)

    paragraphs = document.paragraphs
    cut = 10

    for paragraph in paragraphs:
        if paragraph.text and len(paragraph.text) > 2 and paragraph.text[:2] == "1.":
            break
        lim += 1
    for i in range(0, len(paragraphs) // cut + 1):
        paragraphs_sliced = paragraphs[i * cut + lim : (i + 1) * cut + lim + 1]
        dct = {x: paragraphs_sliced[x].text for x in range(len(paragraphs_sliced))}
        n_dct = {}
        for el, dat in dct.items():
            if dat:
                n_dct[el] = dat
        x = requests.post("http://109.248.175.223:5000/api", json=n_dct)
        try:
            jsn = x.json()
            if x.status_code == 200:
                for j in range(len(paragraphs_sliced)):
                    if j in n_dct:
                        paragraph = paragraphs_sliced[j]
                        el_id, dat = jsn[str(j)]
                        if dat < 50:
                            text = paragraph.text
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                            run.add_text(text)
                            c += 1
            else:
                print("AI server error")
        except InvalidJSONError:
            print("json pars error")
    document.save(file.file.path)
    return pk
